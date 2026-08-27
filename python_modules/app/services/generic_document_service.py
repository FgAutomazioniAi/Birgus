from __future__ import annotations

import json
from dataclasses import dataclass
from io import BytesIO
from typing import Any

from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


@dataclass(frozen=True)
class ContentBlock:
    type: str
    text: str | None = None
    level: int = 1
    items: list[str] | None = None
    rows: list[list[Any]] | None = None


class GenericDocumentService:
    def generate(self, *, content: Any, title: str | None, output_format: str) -> tuple[bytes, str]:
        blocks = self._normalize_content(content)
        if not blocks:
            raise ValueError("Il contenuto risulta vuoto.")

        normalized_format = (output_format or "docx").strip().lower()
        if normalized_format == "docx":
            return self._generate_docx(title, blocks), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if normalized_format == "pdf":
            return self._generate_pdf(title, blocks), "application/pdf"
        if normalized_format in {"md", "markdown"}:
            return self._generate_markdown(title, blocks).encode("utf-8"), "text/markdown; charset=utf-8"

        raise ValueError(f"Formato non supportato: {output_format}")

    def _generate_markdown(self, title: str | None, blocks: list[ContentBlock]) -> str:
        lines: list[str] = []
        if title:
            lines.extend([f"# {title}", ""])

        for block in blocks:
            if block.type == "heading":
                lines.extend([f"{'#' * max(1, min(block.level, 6))} {block.text or ''}".rstrip(), ""])
            elif block.type == "list" and block.items:
                lines.extend([f"- {item}" for item in block.items])
                lines.append("")
            elif block.type == "table" and block.rows:
                rows = [[str(cell) for cell in row] for row in block.rows if isinstance(row, list)]
                if rows:
                    header = rows[0]
                    lines.extend([
                        f"| {' | '.join(header)} |",
                        f"| {' | '.join(['---'] * len(header))} |",
                        *[f"| {' | '.join(row)} |" for row in rows[1:]],
                        "",
                    ])
            elif block.text:
                lines.extend([block.text, ""])

        return "\n".join(lines).rstrip() + "\n"

    def _normalize_content(self, content: Any) -> list[ContentBlock]:
        if isinstance(content, str):
            normalized = content.strip()
            return [ContentBlock(type="paragraph", text=normalized)] if normalized else []

        if isinstance(content, list):
            blocks: list[ContentBlock] = []
            for item in content:
                if isinstance(item, ContentBlock):
                    blocks.append(item)
                elif isinstance(item, dict):
                    blocks.append(ContentBlock(
                        type=str(item.get("type") or "paragraph"),
                        text=self._optional_string(item.get("text")),
                        level=self._positive_int(item.get("level"), 1),
                        items=[str(value) for value in item.get("items", [])] if isinstance(item.get("items"), list) else None,
                        rows=item.get("rows") if isinstance(item.get("rows"), list) else None,
                    ))
            return blocks

        if isinstance(content, dict):
            if "type" in content:
                return self._normalize_content([content])
            return [ContentBlock(type="paragraph", text=json.dumps(content, ensure_ascii=False, indent=2))]

        return []

    def _generate_docx(self, title: str | None, blocks: list[ContentBlock]) -> bytes:
        document = Document()
        if title:
            document.add_heading(title, level=0)

        for block in blocks:
            if block.type == "heading":
                document.add_heading(block.text or "", level=min(max(block.level, 1), 9))
            elif block.type == "list":
                for item in block.items or []:
                    document.add_paragraph(item, style="List Bullet")
            elif block.type == "table" and block.rows:
                rows = block.rows
                table = document.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Light Grid Accent 1"
                for row_index, row in enumerate(rows):
                    for column_index, cell_text in enumerate(row):
                        table.cell(row_index, column_index).text = str(cell_text)
            elif block.text:
                document.add_paragraph(block.text)

        stream = BytesIO()
        document.save(stream)
        return stream.getvalue()

    def _generate_pdf(self, title: str | None, blocks: list[ContentBlock]) -> bytes:
        stream = BytesIO()
        document = SimpleDocTemplate(stream, pagesize=A4)
        styles = getSampleStyleSheet()
        story: list[Any] = []

        if title:
            story.append(Paragraph(title, styles["Title"]))
            story.append(Spacer(1, 16))

        for block in blocks:
            if block.type == "heading":
                style = styles["Heading1"] if block.level <= 1 else styles["Heading2"] if block.level == 2 else styles["Heading3"]
                story.append(Paragraph(block.text or "", style))
                story.append(Spacer(1, 8))
            elif block.type == "list":
                items = [ListItem(Paragraph(item, styles["Normal"])) for item in (block.items or [])]
                story.append(ListFlowable(items, bulletType="bullet"))
                story.append(Spacer(1, 8))
            elif block.type == "table" and block.rows:
                table = Table(block.rows)
                table.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                ]))
                story.append(table)
                story.append(Spacer(1, 8))
            elif block.text:
                story.append(Paragraph(block.text.replace("\n", "<br/>"), styles["Normal"]))
                story.append(Spacer(1, 8))

        document.build(story)
        return stream.getvalue()

    def _optional_string(self, value: Any) -> str | None:
        return value.strip() if isinstance(value, str) and value.strip() else None

    def _positive_int(self, value: Any, fallback: int) -> int:
        try:
            parsed = int(value)
        except (TypeError, ValueError):
            return fallback
        return parsed if parsed > 0 else fallback
