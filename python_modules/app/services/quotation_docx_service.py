from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


class QuotationDocxService:
    _FIELD_LABELS: tuple[tuple[str, str], ...] = (
        ("Title", "Title:"),
        ("Printing/Press", "Printing / Press:"),
        ("Imposition", "Imposition:"),
        ("Trim size", "Trim size:"),
        ("Extent", "Extent:"),
        ("Text", "Text:"),
        ("1st form", "1st form:"),
        ("Endpapers", "Endpapers:"),
        ("Casecover", "Casecover:"),
        ("Dust jacket", "Dust jacket:"),
        ("Binding", "Binding:"),
        ("Packing", "Packing:"),
        ("Cartons", "Cartons:"),
        ("Transport", "Transport:"),
        ("Prices", "Prices:"),
        ("Extra costs", "Extra costs:"),
    )

    def build_quotation_docx(self, structured_data: dict[str, Any]) -> bytes:
        document = Document()
        self._configure_section(document.sections[0])

        self._append_optional_line(document, self._compose_place_date(structured_data), space_after=Pt(6))
        self._append_optional_line(document, self._prefix_if_present("Attn. ", structured_data.get("Attn")))
        self._append_optional_line(document, self._normalized_string(structured_data.get("Company")))
        self._append_optional_line(document, self._normalized_string(structured_data.get("Address1")))
        self._append_optional_line(document, self._normalized_string(structured_data.get("Address2")))
        self._append_spacer(document)
        self._append_optional_line(document, self._prefix_if_present("RE: ", structured_data.get("Reference")), bold=True)
        self._append_spacer(document)
        self._append_optional_line(document, self._normalized_string(structured_data.get("Greeting")))
        self._append_spacer(document)
        self._append_specification_table(document, structured_data)
        self._append_spacer(document)
        self._append_optional_line(document, self._normalized_string(structured_data.get("ClosingHeaderAttn")))
        self._append_optional_line(
            document,
            self._prefix_if_present("RE: ", structured_data.get("ClosingReference")),
            bold=True,
        )
        self._append_spacer(document)
        self._append_optional_paragraph(document, self._normalized_string(structured_data.get("ClosingParagraph1")))
        self._append_optional_paragraph(document, self._normalized_string(structured_data.get("ClosingParagraph2")))
        self._append_spacer(document)
        self._append_optional_line(document, self._normalized_string(structured_data.get("Signoff")))
        self._append_optional_line(document, self._normalized_string(structured_data.get("Signature")))

        stream = BytesIO()
        document.save(stream)
        return stream.getvalue()

    def _configure_section(self, section: Any) -> None:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.footer_distance = Cm(1)
        footer = section.footer
        footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.paragraphs[0].text = (
            "Birgus srl | Headquarter | Via Giuseppe Garibaldi, 5/41 | 37057 San Giovanni Lupatoto | Verona | Italia"
        )
        self._style_paragraph(footer.paragraphs[0], size_pt=6)
        footer_line = footer.add_paragraph("Tel. +39 045 8781396 | info@birgus.com | www.birgus.com")
        footer_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._style_paragraph(footer_line, size_pt=6)

    def _append_specification_table(self, document: Document, structured_data: dict[str, Any]) -> None:
        table = document.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        for field_key, label in self._FIELD_LABELS:
            row = table.add_row()
            label_cell = row.cells[0]
            value_cell = row.cells[1]
            label_cell.width = Cm(4.5)
            value_cell.width = Cm(14.5)
            label_paragraph = label_cell.paragraphs[0]
            label_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            label_run = label_paragraph.add_run(label)
            label_run.bold = True
            label_run.font.name = "Arial"
            label_run.font.size = Pt(9)
            self._set_cell_vertical_alignment(label_cell)

            value_lines = self._split_lines(structured_data.get(field_key))
            for index, line in enumerate(value_lines or [""]):
                paragraph = value_cell.paragraphs[0] if index == 0 else value_cell.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(1)
                run = paragraph.add_run(line)
                run.font.name = "Arial"
                run.font.size = Pt(9)
            self._set_cell_vertical_alignment(value_cell)

    def _append_optional_paragraph(self, document: Document, text: str | None) -> None:
        if not text:
            return

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(10)

    def _append_optional_line(self, document: Document, text: str | None, *, bold: bool = False, space_after: Any = Pt(1)) -> None:
        if not text:
            return

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = space_after
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.name = "Arial"
        run.font.size = Pt(10)

    def _append_spacer(self, document: Document) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run("")
        run.font.name = "Arial"
        run.font.size = Pt(10)

    def _compose_place_date(self, structured_data: dict[str, Any]) -> str:
        place = self._normalized_string(structured_data.get("Place")) or "San Giovanni Lupatoto"
        date = self._normalized_string(structured_data.get("Date"))
        return f"{place}, {date}" if date else place

    def _prefix_if_present(self, prefix: str, value: Any) -> str | None:
        normalized = self._normalized_string(value)
        return f"{prefix}{normalized}" if normalized else None

    def _normalized_string(self, value: Any) -> str | None:
        if not isinstance(value, str):
            return None
        normalized = " ".join(value.strip().split())
        return normalized or None

    def _split_lines(self, value: Any) -> list[str]:
        if not isinstance(value, str):
            return []
        return [line.strip() for line in value.splitlines() if line.strip()]

    def _style_paragraph(self, paragraph: Any, *, size_pt: int) -> None:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(size_pt)

    def _set_cell_vertical_alignment(self, cell: Any) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        v_align = OxmlElement("w:vAlign")
        v_align.set(qn("w:val"), "center")
        tc_pr.append(v_align)
